"""Document extractors — turn binary .xlsx/.docx/.pdf uploads into the *normalized text* of an
existing import format (CSV-text or Markdown-text), so the proven importers stay the single
source of truth. Pure and keyless: no network, no secrets, never log file contents. Extraction
is intentionally fixable — the caller shows the returned text for review before freezing.
"""

from __future__ import annotations

import csv
import io
import re
from typing import Literal

from pydantic import BaseModel

from orionfold.data.importers import ImportFormat

DocFormat = Literal["xlsx", "docx", "pdf"]

_EXT_TO_FORMAT: dict[str, DocFormat] = {".xlsx": "xlsx", ".docx": "docx", ".pdf": "pdf"}


class ExtractResult(BaseModel):
    """Normalized text in an existing import format, plus extraction-level warnings."""

    format: ImportFormat
    text: str
    warnings: list[str]


class DocExtractError(ValueError):
    """A document could not be opened or read — surfaced to the API as HTTP 422."""


def doc_format_for(filename: str) -> DocFormat | None:
    for ext, fmt in _EXT_TO_FORMAT.items():
        if filename.lower().endswith(ext):
            return fmt
    return None


def extract_document(data: bytes, doc_format: DocFormat) -> ExtractResult:
    if doc_format == "xlsx":
        return _extract_xlsx(data)
    if doc_format == "docx":
        return _extract_docx(data)
    if doc_format == "pdf":
        return _extract_pdf(data)
    raise DocExtractError(f"Unsupported document format: {doc_format}")


def _extract_xlsx(data: bytes) -> ExtractResult:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # openpyxl raises various errors on bad input
        raise DocExtractError(f"Could not read the Excel file: {exc}") from exc
    ws = wb.active
    rows = ws.iter_rows(values_only=True) if ws is not None else iter(())
    header = next(rows, None)
    warnings: list[str] = []
    if not header:
        return ExtractResult(format="csv", text="", warnings=["The spreadsheet was empty."])
    cols = {str(c).strip().lower(): i for i, c in enumerate(header) if c is not None}
    in_i = cols.get("input", cols.get("input_text"))
    out_i = cols.get("expected", cols.get("expected_text"))
    if in_i is None or out_i is None:
        warnings.append(
            "Spreadsheet needs an 'input' and an 'expected' column header. "
            "Rename your columns, or edit the text below into input,expected rows."
        )
        return ExtractResult(format="csv", text="", warnings=warnings)
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["input", "expected"])
    for row in rows:
        a = "" if in_i >= len(row) or row[in_i] is None else str(row[in_i]).strip()
        b = "" if out_i >= len(row) or row[out_i] is None else str(row[out_i]).strip()
        if a or b:
            writer.writerow([a, b])
    return ExtractResult(format="csv", text=buf.getvalue(), warnings=warnings)


def _extract_docx(data: bytes) -> ExtractResult:
    try:
        import docx

        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocExtractError(f"Could not read the Word file: {exc}") from exc
    # Prefer two-column tables (input | expected); fall back to paragraph headings.
    blocks: list[str] = []
    for table in doc.tables:
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            if ri == 0 and cells[0].lower() in {"input", "input_text"}:
                continue  # skip header row
            if cells[0] or cells[1]:
                blocks.append(f"## Input\n{cells[0]}\n\n## Expected\n{cells[1]}\n\n---")
    if blocks:
        return ExtractResult(format="markdown", text="\n".join(blocks), warnings=[])
    paras = "\n".join(p.text for p in doc.paragraphs)
    text, warnings = normalize_pairs_to_markdown(paras)
    return ExtractResult(format="markdown", text=text, warnings=warnings)


def _pdf_text(data: bytes) -> str:
    """Isolated so tests can monkeypatch it without a binary fixture."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_pdf(data: bytes) -> ExtractResult:
    try:
        raw = _pdf_text(data)
    except Exception as exc:
        raise DocExtractError(f"Could not read the PDF file: {exc}") from exc
    text, warnings = normalize_pairs_to_markdown(raw)
    warnings.insert(
        0,
        "PDF text extraction is lossy — review the pairs below and fix any split text "
        "before freezing.",
    )
    return ExtractResult(format="markdown", text=text, warnings=warnings)


_BARE_LABEL = re.compile(r"^(input|expected)\s*:?\s*$", re.IGNORECASE)


def normalize_pairs_to_markdown(text: str) -> tuple[str, list[str]]:
    """Promote bare 'Input'/'Expected' lines to Markdown headings the importer understands.
    Lines already starting with '#' are left as-is. Returns (markdown_text, warnings)."""
    out: list[str] = []
    saw_label = False
    for raw in text.splitlines():
        line = raw.rstrip()
        m = _BARE_LABEL.match(line.strip())
        if m:
            saw_label = True
            out.append(f"## {m.group(1).capitalize()}")
        else:
            out.append(line)
    warnings: list[str] = []
    if not saw_label and "#" not in text:
        warnings.append(
            "No 'Input'/'Expected' structure found. Add '## Input' / '## Expected' headings "
            "(separated by '---') so the pairs can be parsed."
        )
    return "\n".join(out), warnings
